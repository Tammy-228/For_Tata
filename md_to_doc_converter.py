import tkinter as tk
from tkinter import filedialog
from pathlib import Path
import mammoth
import markdown
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def select_folder():
    """打开文件夹选择对话框"""
    root = tk.Tk()
    root.withdraw()
    folder_path = filedialog.askdirectory(
        title="选择包含 Markdown 文件的文件夹"
    )
    return folder_path


def convert_md_to_doc(md_file_path, output_dir):
    """将单个 MD 文件转换为 DOC 文件"""
    try:
        # 读取 MD 文件内容
        with open(md_file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()

        # 创建新的 Word 文档
        doc = docx.Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(12)

        # 将 Markdown 转换为 HTML
        html_content = markdown.markdown(md_content)

        # 分割内容为段落
        paragraphs = html_content.split('\n')

        # 处理每个段落
        for para in paragraphs:
            if para.startswith('# '):  # 一级标题
                p = doc.add_paragraph()
                run = p.add_run(para[2:])
                run.font.size = Pt(18)
                run.font.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif para.startswith('## '):  # 二级标题
                p = doc.add_paragraph()
                run = p.add_run(para[3:])
                run.font.size = Pt(16)
                run.font.bold = True
            elif para.startswith('### '):  # 三级标题
                p = doc.add_paragraph()
                run = p.add_run(para[4:])
                run.font.size = Pt(14)
                run.font.bold = True
            elif para.strip():  # 普通段落
                p = doc.add_paragraph()
                run = p.add_run(para)
                run.font.size = Pt(12)

        # 生成输出文件名
        output_filename = md_file_path.stem + '.docx'
        output_path = output_dir / output_filename

        # 保存文档
        doc.save(str(output_path))
        print(f"已转换: {md_file_path.name} -> {output_filename}")

        return True
    except Exception as e:
        print(f"转换失败 {md_file_path.name}: {str(e)}")
        return False


def main():
    # 选择输入文件夹
    input_folder = select_folder()
    if not input_folder:
        print("未选择文件夹")
        return

    # 创建输出目录
    input_path = Path(input_folder)
    output_dir = input_path / "转换后的文档"
    output_dir.mkdir(exist_ok=True)

    # 获取所有 MD 文件
    md_files = list(input_path.glob('*.md'))

    if not md_files:
        print("所选文件夹中没有找到 Markdown 文件")
        return

    print(f"找到 {len(md_files)} 个 Markdown 文件")

    # 转换每个文件
    success_count = 0
    for md_file in md_files:
        if convert_md_to_doc(md_file, output_dir):
            success_count += 1

    # 打印统计信息
    print(f"\n转换完成:")
    print(f"总文件数: {len(md_files)}")
    print(f"成功转换: {success_count}")
    print(f"失败数量: {len(md_files) - success_count}")
    print(f"\n转换后的文件保存在: {output_dir}")


if __name__ == "__main__":
    main()
