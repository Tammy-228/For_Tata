import tkinter as tk
from tkinter import filedialog
from pathlib import Path
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os


def select_folder():
    """打开文件夹选择对话框"""
    root = tk.Tk()
    root.withdraw()
    folder_path = filedialog.askdirectory(
        title="选择包含 Markdown 文件的文件夹"
    )
    return folder_path


def clean_markdown_content(content):
    """清理 Markdown 格式，只保留纯文本"""
    # 移除 HTML 标签
    content = re.sub(r'<[^>]+>', '', content)

    # 移除 Markdown 标题标记 (#)
    content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)

    # 移除粗体标记 (**)
    content = re.sub(r'\*\*(.*?)\*\*', r'\1', content)

    # 移除斜体标记 (*)
    content = re.sub(r'\*(.*?)\*', r'\1', content)

    # 移除链接格式，只保留文本
    content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)

    # 移除图片标记
    content = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', '', content)

    # 移除代码块
    content = re.sub(r'```[\s\S]*?```', '', content)

    # 移除行内代码
    content = re.sub(r'`([^`]+)`', r'\1', content)

    # 移除引用标记
    content = re.sub(r'^\s*>\s*', '', content, flags=re.MULTILINE)

    # 移除列表标记
    content = re.sub(r'^\s*[-*+]\s+', '', content, flags=re.MULTILINE)
    content = re.sub(r'^\s*\d+\.\s+', '', content, flags=re.MULTILINE)

    # 移除多余的空行
    content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)

    return content.strip()


def convert_md_to_doc(md_file_path, output_dir):
    """将单个 MD 文件转换为 DOC 文件"""
    try:
        # 读取 MD 文件内容
        with open(md_file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        # 清理 Markdown 格式
        cleaned_content = clean_markdown_content(content)

        # 按段落分割内容
        paragraphs = cleaned_content.split('\n\n')

        # 创建新的 Word 文档
        doc = docx.Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(12)

        # 添加标题（使用文件名作为标题）
        title = md_file_path.stem.replace('[', '').replace(']', '')
        heading = doc.add_paragraph()
        heading_run = heading.add_run(title)
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加空行
        doc.add_paragraph()

        # 添加正文段落
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                run = p.add_run(para.strip())
                run.font.size = Pt(12)

        # 生成输出文件名
        output_filename = md_file_path.stem + '.docx'
        output_path = output_dir / output_filename

        # 保存文档
        doc.save(str(output_path))
        print(f"已转换: {md_file_path.name} -> {output_filename}")

        return True
    except Exception as e:
        print(f"转换失败 {md_file_path.name}: {str(e)}")
        return False


def main():
    # 选择输入文件夹
    input_folder = select_folder()
    if not input_folder:
        print("未选择文件夹")
        return

    # 创建输出目录
    input_path = Path(input_folder)
    output_dir = input_path / "Word文档"
    output_dir.mkdir(exist_ok=True)

    # 获取所有 MD 文件
    md_files = list(input_path.glob('*.md'))

    if not md_files:
        print("所选文件夹中没有找到 Markdown 文件")
        return

    print(f"找到 {len(md_files)} 个 Markdown 文件")

    # 转换每个文件
    success_count = 0
    for md_file in md_files:
        if convert_md_to_doc(md_file, output_dir):
            success_count += 1

    # 打印统计信息
    print(f"\n转换完成:")
    print(f"总文件数: {len(md_files)}")
    print(f"成功转换: {success_count}")
    print(f"失败数量: {len(md_files) - success_count}")
    print(f"\n转换后的文件保存在: {output_dir}")


if __name__ == "__main__":
    main()
